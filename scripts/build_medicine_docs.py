from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source"
PRIMARY = ROOT / "primary"
SUPP = ROOT / "supplemental"
COVER = ROOT / "cover"

MANUSCRIPT_MD = SOURCE / "manuscript_medicine.md"
COVER_MD = COVER / "Cover_Letter_Medicine.md"
TABLE1_CSV = SOURCE / "table1_threshold_crosswalk.csv"
STROBE_TEMPLATE = SOURCE / "STROBE-MR_template.docx"

TITLE = (
    "Circulating immune-cell traits and cardiac conduction disease: "
    "a phenotype-wide Mendelian randomization study"
)

# AMA-style in-text reference numbers, e.g. [1], [3,4], [12-16], rendered superscript
CITATION_PATTERN = r"\[\d+(?:[,\-]\d+)*\]"


def setup_document(line_numbers: bool = False) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        if not line_numbers:
            continue
        sect_pr = section._sectPr
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "continuous")
        ln.set(qn("w:distance"), "360")
        # CT_SectPr is a sequence: lnNumType must precede pgNumType/cols/docGrid.
        # Appending puts it last, which is schema-invalid and can be dropped by a
        # strict consumer, so insert it ahead of the first element that must follow it.
        successors = ("w:pgNumType", "w:cols", "w:formProt", "w:vAlign", "w:docGrid")
        anchor = next((c for c in sect_pr if c.tag in {qn(t) for t in successors}), None)
        if anchor is None:
            sect_pr.append(ln)
        else:
            anchor.addprevious(ln)
    return doc


def add_runs(paragraph, text: str) -> None:
    text = text.replace(" x 10", " × 10")
    text = text.replace("<=", "≤").replace(">=", "≥")
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = part.startswith("*") and part.endswith("*") and not bold
        inner = part[2:-2] if bold else part[1:-1] if italic else part
        for token in re.split(r"(\^-?\d+|" + CITATION_PATTERN + r")", inner):
            if not token:
                continue
            superscript = token.startswith("^") or re.fullmatch(CITATION_PATTERN, token)
            run = paragraph.add_run(token[1:] if token.startswith("^") else token)
            run.bold = bold
            run.italic = italic
            if superscript:
                run.font.superscript = True


def sci(value: float, digits: int = 2) -> str:
    """Format a P value as 3.06 x 10^-5, the notation add_runs renders and the body text uses.
    Table cells previously carried the raw 3.06e-05 exponent form."""
    mantissa, exponent = f"{value:.{digits}e}".split("e")
    return f"{mantissa} x 10^{int(exponent)}"


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_runs(p, text)
    return p


def split_refs(md: str) -> tuple[str, str]:
    body, refs = md.split("\n## References\n", 1)
    return body.rstrip(), refs.strip()


def section_text(md: str, header: str) -> str:
    match = re.search(rf"## {re.escape(header)}\n\n(.*?)(?=\n## |\Z)", md, re.S)
    return match.group(1).strip() if match else ""


def add_heading(doc: Document, title: str, level: int = 1) -> None:
    # Medicine numbers body sections as "1.", "2.1."; the in-text cross-references
    # ("Section 3.1") only resolve if the numbers survive into the heading.
    clean = re.sub(r"^(\d+(?:\.\d+)*)\s+", r"\1. ", title.strip())
    doc.add_heading(clean, level=level)


TABLE1_TITLE = (
    "Table 1. Genome-wide-threshold signals in relation to the primary screen. "
    "The five threshold-specific immunophenotypes reaching FDR significance for "
    "cardiac conduction disorders under the stricter genome-wide instrument threshold "
    "(P < 5 x 10^-8) are shown against their position in the primary screen "
    "(P < 1 x 10^-5). None reached nominal significance in the primary analysis. "
    "Variant-level composition of these five instrument sets is given in "
    "Supplementary Table S14."
)
TABLE1_ABBREV = (
    "CI, confidence interval; FDR, false discovery rate; GWS, genome-wide significant "
    "(P < 5 x 10^-8); IV, instrumental variable; IVW, inverse-variance weighted; "
    "OR, odds ratio; Wald, single-variant Wald ratio."
)


def add_table1(doc: Document, grid: bool = True) -> None:
    """Render Table 1. The manuscript carries the caption only, because Medicine requires
    each table to be uploaded as its own file; grid=True is used for that separate file."""
    rows = []
    with TABLE1_CSV.open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            if row["outcome"] == "Cardiac conduction disorders":
                rows.append(row)

    add_paragraph(doc, TABLE1_TITLE + (" " + TABLE1_ABBREV if not grid else ""))
    if not grid:
        return
    columns = [
        "Immunophenotype",
        "GWS N IVs",
        "Method",
        "GWS OR (95% CI)",
        "GWS P",
        "GWS FDR",
        "Primary rank",
        "Primary P",
        "Primary FDR",
    ]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        run = table.rows[0].cells[i].paragraphs[0].add_run(col)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        values = [
            row["trait"],
            str(int(float(row["gws_nIV"]))),
            row["gws_method"],
            f"{float(row['gws_OR']):.3f} ({float(row['gws_OR_L']):.3f}-{float(row['gws_OR_U']):.3f})",
            sci(float(row['gws_p'])),
            f"{float(row['gws_FDR']):.3f}".lstrip("0"),
            f"{int(float(row['main_rank']))}/731",
            f"{float(row['main_p']):.2f}".lstrip("0"),
            f"{float(row['main_FDR']):.2f}".lstrip("0"),
        ]
        for i, value in enumerate(values):
            add_runs(cells[i].paragraphs[0], value)
    for table_row in table.rows:
        for cell in table_row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    add_paragraph(doc, TABLE1_ABBREV)


# Medicine asks for figure legends on their own page after the reference list, not beside the
# callout, so add_body collects them here and add_backmatter writes them out.
COLLECTED_LEGENDS: list[str] = []


def add_figure_legend(doc: Document, text: str) -> None:
    COLLECTED_LEGENDS.append(text)


def add_frontmatter(doc: Document, md: str) -> None:
    title = md.splitlines()[0].lstrip("# ").strip()
    p = add_paragraph(doc, title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(15)

    p = add_paragraph(doc, "Wei Li, Yanfu Wang*, Fang Liu and Yuanzheng Li")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "Department of Cardiovascular Medicine, Aviation General Hospital, Beijing 100012, China")
    add_paragraph(doc, "*Correspondence: Yanfu Wang, Department of Cardiovascular Medicine, Aviation General Hospital, Beijing 100012, China; Email: wangfy328@email.hkzyy.com.cn")

    abbreviations = section_text(md, "Abbreviations")
    if abbreviations:
        add_heading(doc, "Abbreviations")
        add_paragraph(doc, abbreviations.strip())

    abstract = re.search(r"## Abstract\n\n(.*?)\n\n\*\*Keywords", md, re.S).group(1).strip()
    add_heading(doc, "Abstract")
    for paragraph in re.split(r"\n\s*\n", abstract):
        add_paragraph(doc, paragraph.strip())
    keywords = re.search(r"\*\*Keywords:\*\*\s*(.*)", md).group(1).strip()
    add_paragraph(doc, "Keywords: " + keywords)


# Display items are inserted by matching a trigger phrase in the manuscript text. If the text
# is reworded the match silently stops firing and the item vanishes from the DOCX, so every
# trigger is registered here and verified after the body is written.
DISPLAY_TRIGGERS = {
    "Figure 1": "The workflow is summarized in Figure 1",
    "Figure 2": "found no FDR-significant association for either outcome (Figure 2)",
    "Table 1": "primary-screen result (Table 1)",
    "Figure 3": "(Figure 3B)",
}


def add_body(doc: Document, md: str) -> None:
    emitted: set[str] = set()
    before_refs, _ = split_refs(md)
    body = before_refs.split("## 1 Introduction", 1)[1].split("## Author contributions", 1)[0]
    for raw_line in ("## 1 Introduction\n" + body).splitlines():
        line = raw_line.rstrip()
        if not line or line == "---":
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)", line)
        if heading:
            level = max(1, min(3, len(heading.group(1)) - 1))
            add_heading(doc, heading.group(2), level)
            continue
        add_paragraph(doc, line)
        if DISPLAY_TRIGGERS["Figure 1"] in line:
            emitted.add("Figure 1")
            add_figure_legend(
                doc,
                "Figure 1. Study flowchart of the bidirectional Mendelian randomization analysis, "
                "including the genome-wide instrument-threshold sensitivity analysis.",
            )
        if DISPLAY_TRIGGERS["Figure 2"] in line:
            emitted.add("Figure 2")
            add_figure_legend(
                doc,
                "Figure 2. Volcano plots from the primary P < 1 x 10^-5 instrument-threshold analysis of 731 immune-cell phenotypes and cardiac conduction disorders or atrioventricular block. No phenotype survived false-discovery-rate correction in the primary screen. Green labels identify five nominally associated phenotypes that met the descriptive sensitivity-stability criteria defined in Section 2.5; none was FDR-significant. Gold diamonds mark the genome-wide-threshold FDR signals, the five phenotype-level signals identified only under the stricter P < 5 x 10^-8 threshold. The two sets are non-overlapping and share no phenotype. The genome-wide-threshold FDR signals are plotted in both panels to show where they fall in the primary screen; they reached FDR significance for cardiac conduction disorders only.",
            )
        if DISPLAY_TRIGGERS["Table 1"] in line:
            emitted.add("Table 1")
            add_table1(doc, grid=False)
        if DISPLAY_TRIGGERS["Figure 3"] in line:
            emitted.add("Figure 3")
            add_figure_legend(
                doc,
                "Figure 3. Threshold dependence of the phenotype-wide results for cardiac conduction disorders. (A) Rank-rank plot of the phenotype-wide P-value ordering under the primary (P < 1 x 10^-5) and genome-wide (P < 5 x 10^-8) instrument thresholds, across the 607 phenotype records analyzable under both (Spearman rho = .11). Gold points mark the genome-wide-threshold FDR signals, none of which was even nominally significant in the primary screen. The LOESS curve is shown for visual orientation only. (B) Distribution of the number of instrumental variants per phenotype under the two thresholds, shown as a percentage of analyzable phenotypes because the denominators differ (607 at the genome-wide threshold, 731 at the primary threshold). At the genome-wide threshold, 26.5% of phenotypes were instrumented by a single variant and 68.0% by three or fewer; the genome-wide-threshold FDR signals rested on one to four variants. At the primary threshold, no phenotype was instrumented by a single variant, only 3 of 731 had three or fewer, and the median instrument count was 22. Post-hoc diagnostics for those signals, including the share of inverse-variance weight carried by a single rare variant and the effect of excluding instruments with a minor allele frequency below 1%, are given in Supplementary Tables S14 and S15.",
            )
    missing = sorted(set(DISPLAY_TRIGGERS) - emitted)
    if missing:
        raise RuntimeError(
            "display items were not inserted because their trigger phrase no longer appears "
            f"in the manuscript: {', '.join(missing)}. Update DISPLAY_TRIGGERS to match the text."
        )


# One entry per uploaded file, in the order the files are cited, as the journal asks for a
# separate set of SDC legends after the figure legends. The 23 supplementary tables travel in a
# single workbook, so they are one numbered item rather than 23.
SDC_LEGENDS = [
    "Supplemental Digital Content 1. Supplementary Tables S1-S23 (Excel workbook, one table per "
    "sheet): full MR results for all 731 immunophenotypes on each outcome (S1, S2); cross-outcome "
    "concordant phenotypes (S3); sensitivity analyses (S4); reverse MR (S5); genetic instruments "
    "(S6); leave-one-out analyses (S7); MR-PRESSO and Steiger directionality (S8); power (S9, S10, "
    "S18); the genome-wide instrument-threshold sensitivity analysis and cross-threshold comparison "
    "(S11, S12); the targeted lymphocyte-count re-analysis (S13); variant-level composition and "
    "post-hoc diagnostics for the genome-wide-threshold findings (S14-S16); outcome endpoint "
    "definitions (S17); diagnostic computability and Cochran's Q (S19, S20); the primary screen "
    "under a per-instrument credibility filter (S21); the atrial-fibrillation specificity control "
    "(S22); and the bundle-branch-block component-endpoint screen (S23).",
    "Supplemental Digital Content 2. Supplementary Figure S1 (image): forest plot of the "
    "cross-outcome concordant immunophenotypes.",
    "Supplemental Digital Content 3. Supplementary Figure S2 (image): leave-one-out plots for the "
    "sensitivity-stable phenotypes.",
    "Supplemental Digital Content 4. Completed STROBE-MR reporting checklist (Word document).",
]


def add_backmatter(doc: Document, md: str) -> None:
    headers = [
        "Author contributions",
        "Acknowledgments",
        "Funding",
        "Ethics approval",
        "Conflicts of interest",
        "Data availability",
        "Supplemental Digital Content",
    ]
    for header in headers:
        content = section_text(md, header)
        if content:
            add_heading(doc, header)
            # Author contributions lists one CRediT role per line; keep the breaks.
            for line in content.splitlines():
                if line.strip():
                    add_paragraph(doc, line.strip())

    _, refs = split_refs(md)
    add_heading(doc, "References")
    for line in refs.splitlines():
        if line.strip():
            add_paragraph(doc, line.strip())

    doc.add_page_break()
    add_heading(doc, "Figure legends")
    for legend in COLLECTED_LEGENDS:
        p = add_paragraph(doc, legend)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_heading(doc, "Supplemental Digital Content legends")
    for item in SDC_LEGENDS:
        add_paragraph(doc, item)


def build_manuscript() -> None:
    md = MANUSCRIPT_MD.read_text(encoding="utf-8")
    doc = setup_document(line_numbers=True)
    add_frontmatter(doc, md)
    add_body(doc, md)
    add_backmatter(doc, md)
    PRIMARY.mkdir(parents=True, exist_ok=True)
    doc.save(PRIMARY / "Manuscript_Medicine.docx")


def build_table1() -> None:
    doc = setup_document()
    add_table1(doc, grid=True)
    PRIMARY.mkdir(parents=True, exist_ok=True)
    doc.save(PRIMARY / "Table_1.docx")


def build_cover_letter() -> None:
    md = COVER_MD.read_text(encoding="utf-8")
    doc = setup_document()
    # Keep the letter on one page. The default template adds 10pt after every
    # paragraph at 1.15 line spacing, and blank source lines were emitted as
    # empty paragraphs carrying the same spacing, which together pushed the
    # letter onto a second page. Paragraph separation now comes from space_after.
    normal = doc.styles["Normal"].paragraph_format
    normal.space_after = Pt(5)
    normal.line_spacing = 1.0
    for line in md.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            p = add_paragraph(doc, line[2:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
            continue
        add_paragraph(doc, line.rstrip())
    doc.save(COVER / "Cover_Letter_Medicine.docx")


# row index in the official template -> (section pointers, relevant-text cell)
STROBE_FILL = {1: (['Title', 'Abstract'], 'Title; Abstract (title and Methods paragraph)'), 3: (['Introduction'], 'Introduction, paragraphs 1-3'), 4: (['Introduction'], 'Introduction, paragraphs 4-6'), 6: (['Methods 2.1', '2.2', '2.3', '2.4', 'Figure 1'], 'Key design elements are presented in Methods 2.1-2.4 and summarized in Figure 1; the data sources for every phase are listed in Methods 2.2 and 2.3, with the outcome endpoint definitions tabulated in Supplementary Table S17. Sub-items a-e follow.'), 7: (['Methods 2.1', '2.2', '2.3'], '4a: Methods 2.1. Relevant dates for data collection are given in Methods 2.2 (GWAS Catalog accessions retrieved through OpenGWAS in July 2026) and Methods 2.3 (FinnGen release 11 downloaded on 16 July 2026); recruitment and follow-up periods of the two source cohorts are not restated in this article.'), 8: (['Methods 2.2', '2.3', '2.6', 'Results 3.1'], '4b: Methods 2.2 and 2.3 (exposure cohort of 3,757 Sardinian participants with accession-level sample sizes of 1,244 to 3,669, and FinnGen release 11 case and control counts of 12,371/342,690 for cardiac conduction disorders and 6,935/342,690 for atrioventricular block; outcome case and control definitions, including the contributing diagnostic codes, are referenced to the FinnGen endpoint browser and summarized in Supplementary Table S17). Power is reported in Methods 2.6 and Results 3.1 and in Supplementary Tables S9, S10 and S18 as a selection-conditioned upper bound rather than achieved power; the manuscript does not state that a power or sample size calculation was carried out prior to the main analysis.'), 9: (['Methods 2.4'], '4c: Methods 2.4.'), 10: (['Methods 2.2', '2.3'], '4d: Methods 2.2, 2.3 (outcome case and control definitions, including the contributing diagnostic codes, are referenced to the FinnGen endpoint browser and summarized in Supplementary Table S17).'), 11: (['Methods 2.1', 'Ethics approval'], '4e: Methods 2.1 and the Ethics approval statement.'), 12: (['Methods 2.1'], 'Methods 2.1 (relevance, independence and exclusion-restriction assumptions, followed by the further identifying conditions - linearity and effect homogeneity, or a local interpretation in their absence - under which the ratio estimate recovers a causal effect)'), 13: (['Methods 2.2', '2.4', '2.5', '2.6', '2.8'], 'Statistical methods are described in Methods 2.5 (estimators and diagnostics), 2.6 (multiple testing, power and the instrument-threshold sensitivity analysis) and 2.8 (post-hoc rare-variant diagnostics), with variable handling in Methods 2.2 and 2.4. Sub-items a-e follow.'), 14: (['Methods 2.2'], '6a: Methods 2.2 (traits inverse-normal transformed in the source GWAS; estimates per 1 standard deviation).'), 15: (['Methods 2.4', '2.8'], '6b: Methods 2.4, and Methods 2.8 for the post-hoc minor-allele-frequency filter.'), 16: (['Methods 2.2', '2.5'], '6c: Methods 2.5. Covariate adjustment on the exposure side (sex, age and age squared, applied in the source GWAS) is stated in Methods 2.2; the covariate set used in the FinnGen outcome GWAS is not restated in this article, so the two samples are not documented as sharing an identical adjustment set.'), 17: (['Methods 2.4'], '6d: Methods 2.4 (variants absent from the outcome dataset were dropped; no linkage-disequilibrium proxies were substituted).'), 18: (['Methods 2.6'], '6e: Methods 2.6.'), 19: (['Methods 2.4', '2.5', '2.6', '2.8'], "Methods 2.4 (F-statistic filtering, winner's curse), 2.5 (MR-Egger intercept, Cochran's Q, MR-PRESSO, leave-one-out; MR-PRESSO computed with the MRPRESSO package and used as a global heterogeneity diagnostic only), 2.6 (Steiger directionality), 2.8 (post-hoc variant-level diagnostics)"), 20: (['Methods 2.5', '2.6', '2.7', '2.8'], 'Methods 2.5, 2.6 (reverse-direction MR and multiple testing), 2.7 (targeted re-analysis), 2.8 (post-hoc rare-variant diagnostics and minor-allele-frequency sensitivity analysis), the per-instrument credibility filter, component sub-endpoint screen and atrial-fibrillation negative control described in Methods 2.6 and reported in Supplementary Tables S21, S23 and S22, and the genome-wide instrument-threshold sensitivity analysis in Methods 2.6'), 22: (['Methods 2.4', '2.5', '2.6', 'Data availability'], "9a: Methods 2.6 states that analyses used R 4.6.1, the MendelianRandomization package and ieugwasr. Settings: Methods 2.5 - the primary estimate used inverse-variance weighting from the MendelianRandomization package, which fits fixed-effect IVW with fewer than four variants and random-effects IVW otherwise, with the Wald ratio for single-variant analyses; MR-PRESSO was computed with the MRPRESSO package using 3,000 parametric simulations; Cochran's Q used first-order weights, with the modified second-order weights of Bowden et al examined in sensitivity analysis. Methods 2.4 - linkage-disequilibrium clumping at r2 < .001 within 10 Mb using the 1000 Genomes European reference panel, and exclusion of instruments with F <= 10. Methods 2.6 - Benjamini-Hochberg false-discovery-rate control applied separately within each outcome across 731 phenotypes, and study-wide power at alpha = .05/731. Package-version details are recorded in the session information in the repository given under Data availability."), 23: (['Methods 2.1', '2.8', 'Discussion'], '9b: Methods 2.1 states that the analysis was not prospectively registered. Methods 2.8 and the sixth limitation identify the rare-variant diagnostics as post-hoc. This completed STROBE-MR checklist is provided as Supplemental Digital Content.'), 26: (['Figure 1', 'Methods 2.4', 'Results 3.1', '3.2'], '10a: Figure 1 (study workflow); Methods 2.4 and Results 3.1-3.2 (instrument and phenotype counts retained at each threshold and reasons for exclusion, including the 124 phenotypes not estimable at the genome-wide threshold); Supplementary Tables S11 and S12.'), 27: (['Methods 2.2', '2.3', 'Results 3.1', '3.2'], '10b: Methods 2.2 and 2.3 (exposure trait classes and sample sizes; outcome case and control counts); Results 3.1 and 3.2 (instrument counts and their distributions under each threshold); Supplementary Tables S6 and S18. Only GWAS summary statistics were used, so individual-level phenotypic means, standard deviations and proportions are not available; exposure traits were inverse-normal transformed in the source GWAS (Methods 2.2).'), 28: (['Methods 2.1', '2.5', 'Results 3.4'], '10c: Methods 2.1 (neither source dataset is a meta-analysis of separately published studies, so no across-study heterogeneity assessment applies; instrument-level heterogeneity is reported in Methods 2.5 and Results 3.4).'), 29: (['Methods 2.1', '2.2', '2.3', '2.7'], '10d(i): Methods 2.1, 2.2 and 2.3 (the exposure and outcome samples are drawn from two different European founder populations, Sardinians and Finns; because Sardinia and Finland have population-specific allele frequencies and linkage-disequilibrium patterns, cross-population transportability of the variant-exposure associations is treated as an interpretive caveat rather than demonstrated similarity). 10d(ii): Methods 2.1 states that there is no participant overlap between the exposure and outcome samples. Methods 2.7 notes separately that FinnGen release 11 includes the release 2 participants of the previously reported lymphocyte-count analysis, so that targeted comparison is a re-analysis in an expanded release rather than an independent replication.'), 31: (['Results 3.3', 'Supplemental Digital Content'], '11a: Supplementary Table S6; variant-level variant-exposure and variant-outcome estimates for the five genome-wide-threshold signals are additionally reported in Results 3.3 and Supplementary Table S14.'), 32: (['Results 3.1-3.4', 'Table 1'], '11b: Results 3.1-3.4; Table 1; Supplementary Tables S1, S2, S11, S13.'), 33: ([], '11c: not applicable, because the primary screen returned no false-discovery-rate-significant association and therefore no relative risk to translate into absolute risk.'), 34: (['Figure 2', 'Figure 3'], '11d: Figures 2 and 3; Supplementary Figures S1 and S2.'), 36: (['Results 3.1-3.4'], "12a: Results 3.1, 3.2, 3.3, 3.4; Results 3.1 (phenotype-wide Cochran's Q); Supplementary Tables S4, S8, S14, S19 and S20."), 37: (['Methods 2.5', 'Results 3.1-3.4'], "12b: Results 3.1 (phenotype-wide Cochran's Q under first-order and modified second-order weights), 3.2, 3.3, 3.4 (Cochran's Q, MR-Egger intercept); Methods 2.5 (I-squared GX); Supplementary Tables S4, S8, S14, S19 and S20."), 39: (['Results 3.1-3.4'], '13a: Results 3.1, 3.2, 3.3, 3.4; Supplementary Tables S4, S11, S12, S14, S15, S21 and S23.'), 40: (['Results 3.2-3.4'], '13b: Results 3.2-3.4.'), 41: (['Results 3.3', '3.4'], '13c: Results 3.4 (reverse-direction MR and Steiger filtering); the scope of that analysis is stated at the end of Results 3.3.'), 42: (['Introduction'], '13d: observational associations are summarized in the Introduction, but no formal quantitative comparison with non-MR estimates was performed.'), 43: (['Methods 2.5', 'Results 3.3', '3.4'], '13e: Supplementary Figure S2 (leave-one-out plots) and Supplementary Figure S1 (forest plot); leave-one-out estimates are tabulated in Supplementary Table S7 (Methods 2.5), and the leave-one-out results for the genome-wide-threshold signals are reported in Results 3.3 and the stability of leave-one-out estimates for the five sensitivity-stable phenotypes in Results 3.4.'), 45: (['Discussion', 'Conclusion'], 'Discussion, paragraphs 1-2; Conclusion'), 46: (['Methods 2.1', '2.4', 'Results 3.1-3.3', 'Discussion'], "Discussion, paragraph 10 (limitations First to Sixth); Results 3.1 and Discussion paragraph 3 (power and detectable effect sizes); Supplementary Table S18 (phenotype-level instrument strength and power). Other potential sources of bias: Methods 2.1 (no participant overlap between samples) and 2.4 (winner's curse, which would make the confidence intervals too narrow); Results 3.2 and 3.3 and Supplementary Tables S14-S15 (weight carried by a single rare-variant instrument)."), 48: (['Discussion'], '16a: Discussion, paragraphs 2-10.'), 49: (['Discussion'], '16b: Discussion, paragraphs 5 and 8 (tissue-local and stage-specific mechanisms; gene-environment equivalence).'), 50: (['Discussion', 'Conclusion'], '16c: Discussion, paragraphs 1 and 8 (clinical prioritization; gene-intervention equivalence); Conclusion.'), 51: (['Methods 2.1', 'Discussion'], 'Methods 2.1; Discussion, paragraph 10 (third limitation, founder-population transportability) and the closing sentence of paragraph 10'), 53: (['Funding'], 'Funding'), 54: (['Data availability'], 'Data availability (analysis code and session information at the GitHub repository; immune-cell and lymphocyte-count GWAS summary statistics at the GWAS Catalog; outcome summary statistics at FinnGen release 11); Supplemental Digital Content (Supplementary Tables S1-S23 and Supplementary Figures S1-S2)'), 55: (['Conflicts of interest'], 'Conflicts of interest')}

STROBE_PAGE_NUMBERS = {
    1: "p. 1",
    3: "p. 1",
    4: "p. 1",
    6: "pp. 2-4",
    7: "pp. 2-3",
    8: "pp. 3, 5-6",
    9: "p. 4",
    10: "p. 3",
    11: "pp. 2, 12",
    12: "p. 2",
    13: "pp. 3-6",
    14: "p. 3",
    15: "pp. 4, 6",
    16: "pp. 3-4",
    17: "p. 4",
    18: "p. 5",
    19: "pp. 4-6",
    20: "pp. 4-6",
    22: "pp. 4-5, 12",
    23: "pp. 2, 6, 11",
    26: "pp. 2-7",
    27: "pp. 3, 6-7",
    28: "pp. 2, 4, 9",
    29: "pp. 2-3, 5",
    31: "pp. 8, 13",
    32: "pp. 6-9",
    33: "p. 6",
    34: "pp. 6, 8",
    36: "pp. 6-9",
    37: "pp. 4, 6-9",
    39: "pp. 6-9",
    40: "pp. 7-9",
    41: "pp. 8-9",
    42: "p. 1",
    43: "pp. 4, 8-9",
    45: "pp. 9-11",
    46: "pp. 2, 4, 6-11",
    48: "pp. 9-11",
    49: "pp. 10-11",
    50: "pp. 9, 11",
    51: "pp. 2, 11",
    53: "p. 12",
    54: "pp. 12-13",
    55: "p. 12",
}


def build_strobe() -> None:
    """Fill the journal's own STROBE-MR template rather than a home-made table.

    Medicine supplies a five-column template (Item No. / Section / Checklist item /
    Page No. / Relevant text from manuscript) on Editorial Manager, and it is CC BY 3.0
    material from the Equator Network, so the template file is used as the base document
    and only the two author columns are written. That also keeps the official item
    numbering: our earlier home-made table numbered the items 1-20 without the official
    item 9 (software and pre-registration), so everything from item 9 onward was offset
    by one against the published checklist.

    The Page No. column is populated from the final LibreOffice-rendered manuscript
    PDF used for submission checks. If the manuscript text or layout changes, regenerate
    the PDF and update STROBE_PAGE_NUMBERS before rebuilding this checklist.
    """
    doc = Document(str(STROBE_TEMPLATE))
    table = doc.tables[0]
    if len(table.rows) != 56 or len(table.columns) != 5:
        raise RuntimeError(
            f"STROBE-MR template changed shape ({len(table.rows)}x{len(table.columns)}, "
            "expected 56x5); re-map STROBE_FILL against the new template before building."
        )
    for row_index, (sections, text) in STROBE_FILL.items():
        cells = table.rows[row_index].cells
        cells[3].text = STROBE_PAGE_NUMBERS.get(row_index, "; ".join(sections))
        cells[4].text = text
        for cell in (cells[3], cells[4]):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    SUPP.mkdir(parents=True, exist_ok=True)
    doc.save(SUPP / "STROBE-MR_checklist.docx")


def main() -> None:
    build_manuscript()
    build_table1()
    build_cover_letter()
    build_strobe()
    print("Generated Medicine submission documents in", ROOT)


if __name__ == "__main__":
    main()
